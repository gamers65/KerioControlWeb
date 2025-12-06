# ioc_service.py

from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from starlette.middleware.cors import CORSMiddleware

from PyPDF2 import PdfReader
from docx import Document
import csv
import io
import re
from typing import List

app = FastAPI()

# CORS (разрешить .NET приложению подключаться)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------------
#           IOC EXTRACTOR
# -------------------------------

IP_REGEX = r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
DOMAIN_REGEX = r"\b(?:[a-zA-Z0-9\-]+\.)+[a-zA-Z]{2,}\b"


def normalize_text(txt: str, is_pdf=False) -> str:
    """
    PDF рвет домены и IP -> нужна агрессивная нормализация.
    DOCX/TXT/CSV — нормальные, нельзя удалять все пробелы.
    """

    if is_pdf:
        # для PDF убираем все пробелы и переносы
        txt = re.sub(r"\s+", "", txt)
    else:
        # для DOCX/TXT/CSV — мягкая нормализация
        txt = txt.replace(" ", "")
        txt = txt.replace("\t", "")
        txt = txt.replace("\u00A0", "")  # NBSP

    # деобфускация
    txt = txt.replace("[.]", ".")
    txt = txt.replace("hxxp[:]", "http://")
    txt = txt.replace("hxxps[:]", "https://")

    return txt



def extract_iocs(text: str) -> List[str]:
    ips = re.findall(IP_REGEX, text)
    domains = [d.lower() for d in re.findall(DOMAIN_REGEX, text)]

    all_iocs = sorted(set(ips + domains))

    # удалить мусор типа "09."
    all_iocs = [i for i in all_iocs if len(i) > 3]

    return all_iocs


# -------------------------------
#           PDF
# -------------------------------

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    text = ""

    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"

    return text


# -------------------------------
#           DOCX
# -------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Чтение текста из Word + таблиц (где обычно лежат IOC)."""

    with io.BytesIO(file_bytes) as buffer:
        doc = Document(buffer)

    parts = []

    # 1) обычный текст
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # 2) текст таблиц (тут почти всегда лежат IOC)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    parts.append(value)

    return "\n".join(parts)


# -------------------------------
#           TXT
# -------------------------------

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except:
        return file_bytes.decode("cp1251", errors="ignore")


# -------------------------------
#           CSV
# -------------------------------

def extract_text_from_csv(file_bytes: bytes) -> str:
    buffer = io.StringIO(file_bytes.decode("utf-8", errors="ignore"))
    reader = csv.reader(buffer)

    lines = []
    for row in reader:
        for col in row:
            if col.strip():
                lines.append(col)

    return "\n".join(lines)


# -------------------------------
#          ENDPOINT
# -------------------------------

@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    try:
        content = await file.read()
        filename = file.filename.lower()   # ← ВАЖНО!

        # PDF
        if filename.endswith(".pdf"):
            raw = extract_text_from_pdf(content)
            normalized = normalize_text(raw, is_pdf=True)

        # DOCX
        elif filename.endswith(".docx"):
            raw = extract_text_from_docx(content)
            normalized = normalize_text(raw, is_pdf=False)

        # TXT
        elif filename.endswith(".txt"):
            raw = extract_text_from_txt(content)
            normalized = normalize_text(raw, is_pdf=False)

        # CSV
        elif filename.endswith(".csv"):
            raw = extract_text_from_csv(content)
            normalized = normalize_text(raw, is_pdf=False)

        else:
            return JSONResponse({
                "success": False,
                "message": "Unsupported format"
            })

        iocs = extract_iocs(normalized)

        return JSONResponse({
            "success": True,
            "count": len(iocs),
            "data": iocs
        })

    except Exception as ex:
        return JSONResponse({
            "success": False,
            "message": str(ex)
        })
